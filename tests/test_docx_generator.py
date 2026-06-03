"""Tests for .docx report generation."""
from io import BytesIO
from docx import Document
from src.reports.docx_generator import generate_report_bytes


def test_generate_report_returns_bytes():
    data = {
        "problem": "Build a legal AI platform",
        "solution": [
            {"step": 1, "action": "Set up API", "technology": "FastAPI",
             "skill_strength_score": 5.50},
        ],
        "resources": [
            {"user_id": "U01", "full_name": "Alice", "fit_score": 0.9,
             "priority": "P1", "availability": 30},
        ],
        "risk_flags": [],
    }
    result = generate_report_bytes(data)
    assert isinstance(result, bytes)
    assert len(result) > 100


def test_generate_report_with_risk_flags():
    data = {
        "problem": "Test",
        "solution": [],
        "resources": [
            {"user_id": "U03", "full_name": "Carol", "fit_score": 0.3,
             "priority": "P3", "availability": 5},
        ],
        "risk_flags": [{"user_id": "U03", "reason": "Low fit score: 0.3"}],
    }
    result = generate_report_bytes(data)
    assert isinstance(result, bytes)


def test_generate_report_empty_data():
    data = {
        "problem": "",
        "solution": [],
        "resources": [],
        "risk_flags": [],
    }
    result = generate_report_bytes(data)
    assert isinstance(result, bytes)
    assert len(result) > 100


def test_generate_report_includes_summary_and_tables():
    data = {
        "problem": "Build a conversational insights app",
        "solution": [
            {"step": 1, "action": "Design agents",
             "technology": "LangGraph", "skill_strength_score": 3.5,
             "effort": "2 days"},
        ],
        "task_recommendations": [
            {
                "task_id": 1,
                "task": "Design agents",
                "required_skills": ["LangGraph", "Python"],
                "strength_score": 3.5,
                "resources": [
                    {"user_id": "U01", "full_name": "Alice",
                     "fit_score": 0.9, "availability": 30,
                     "recommendation": "Primary",
                     "matched_skills": ["Python", "FastAPI"]},
                ],
                "rationale": "Selected Alice based on task fit.",
            },
        ],
        "resources": [],
        "risk_flags": [],
    }
    doc = Document(BytesIO(generate_report_bytes(data)))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Executive Summary" in text
    assert "Build a conversational insights app" in text
    assert "Original Query" in text
    assert "How To Read This Report" in text
    assert "Strength:" in text
    assert "Fit:" in text
    assert len(doc.tables) == 2
    table_text = "\n".join(
        cell.text
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert "Design agents" in table_text
    assert "Alice" in table_text
    assert "Resources" in table_text
    assert "P3" not in table_text
    assert "P4" not in table_text
