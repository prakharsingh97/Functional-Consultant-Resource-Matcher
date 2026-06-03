"""Generate a .docx resource matching report with full language support."""
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor

# ── Report label translations ────────────────────────────────────────────────
_LABELS: dict[str, dict[str, str]] = {
    "English": {
        "report_title": "Resource Matching Report",
        "original_query": "Original Query",
        "executive_summary": "Executive Summary",
        "how_to_read": "How To Read This Report",
        "project_breakdown": "Project Breakdown",
        "task_recs": "Task Resource Recommendations",
        "risk_flags_title": "Risk Flags",
        "task": "Task", "workstream": "Workstream",
        "skills_needed": "Skills Needed", "strength": "Strength",
        "effort": "Effort", "resources": "Resources",
        "fit": "Fit", "availability": "Availability", "rationale": "Rationale",
        "hrs_wk": "hrs/wk", "primary": "Primary", "support": "Support",
        "no_tasks": "No project tasks generated.",
        "no_recs": "No task recommendations generated.",
        "needs_staffing": "Needs staffing review",
        "no_risks": "No risk flags.",
        "risk_prefix": "Risk",
        "no_problem": "No problem statement provided.",
        "strength_def": (
            "Strength: aggregate team capability for a skill or task, based on "
            "consultant skill level weights and semantic similarity to the task's "
            "required skills."
        ),
        "fit_def": (
            "Fit: task-level match between a consultant's skills, semantic skill "
            "similarity, and availability. Fit is used only for people recommended "
            "for that task."
        ),
        "summary_template": (
            "Generated a task-based staffing plan with {tasks} project tasks and "
            "{recs} task-level resource recommendations."
        ),
        "risk_summary_template": "{n} task coverage risks need review.",
    },
    "Hindi": {
        "report_title": "संसाधन मिलान रिपोर्ट",
        "original_query": "मूल प्रश्न",
        "executive_summary": "कार्यकारी सारांश",
        "how_to_read": "इस रिपोर्ट को कैसे पढ़ें",
        "project_breakdown": "परियोजना विवरण",
        "task_recs": "कार्य संसाधन अनुशंसाएं",
        "risk_flags_title": "जोखिम संकेत",
        "task": "कार्य", "workstream": "कार्यधारा",
        "skills_needed": "आवश्यक कौशल", "strength": "शक्ति",
        "effort": "प्रयास", "resources": "संसाधन",
        "fit": "फिट", "availability": "उपलब्धता", "rationale": "तर्क",
        "hrs_wk": "घंटे/सप्ताह", "primary": "प्राथमिक", "support": "सहायक",
        "no_tasks": "कोई परियोजना कार्य नहीं बनाया गया।",
        "no_recs": "कोई कार्य अनुशंसाएं नहीं बनाई गईं।",
        "needs_staffing": "स्टाफिंग समीक्षा आवश्यक",
        "no_risks": "कोई जोखिम संकेत नहीं।",
        "risk_prefix": "जोखिम",
        "no_problem": "कोई समस्या विवरण प्रदान नहीं किया गया।",
        "strength_def": (
            "शक्ति: किसी कौशल या कार्य के लिए टीम की समग्र क्षमता, परामर्शदाता "
            "कौशल स्तर के भार और कार्य के आवश्यक कौशल के अर्थपूर्ण समानता के आधार पर।"
        ),
        "fit_def": (
            "फिट: किसी परामर्शदाता के कौशल, अर्थपूर्ण कौशल समानता और उपलब्धता के "
            "बीच कार्य-स्तरीय मिलान।"
        ),
        "summary_template": (
            "{tasks} परियोजना कार्यों और {recs} कार्य-स्तरीय संसाधन अनुशंसाओं "
            "के साथ एक कार्य-आधारित स्टाफिंग योजना तैयार की गई।"
        ),
        "risk_summary_template": "{n} कार्य कवरेज जोखिमों की समीक्षा आवश्यक है।",
    },
    "French": {
        "report_title": "Rapport de Correspondance des Ressources",
        "original_query": "Requête Originale",
        "executive_summary": "Résumé Exécutif",
        "how_to_read": "Comment Lire Ce Rapport",
        "project_breakdown": "Décomposition du Projet",
        "task_recs": "Recommandations de Ressources par Tâche",
        "risk_flags_title": "Indicateurs de Risque",
        "task": "Tâche", "workstream": "Flux de Travail",
        "skills_needed": "Compétences Requises", "strength": "Force",
        "effort": "Effort", "resources": "Ressources",
        "fit": "Adéquation", "availability": "Disponibilité", "rationale": "Justification",
        "hrs_wk": "hrs/sem", "primary": "Primaire", "support": "Support",
        "no_tasks": "Aucune tâche de projet générée.",
        "no_recs": "Aucune recommandation de tâche générée.",
        "needs_staffing": "Révision de dotation nécessaire",
        "no_risks": "Aucun indicateur de risque.",
        "risk_prefix": "Risque",
        "no_problem": "Aucun énoncé de problème fourni.",
        "strength_def": (
            "Force : capacité globale de l'équipe pour une compétence ou une tâche, "
            "basée sur les poids de niveau de compétence des consultants et la "
            "similarité sémantique avec les compétences requises."
        ),
        "fit_def": (
            "Adéquation : correspondance au niveau de la tâche entre les compétences "
            "d'un consultant, la similarité sémantique et la disponibilité."
        ),
        "summary_template": (
            "Plan de dotation basé sur les tâches généré avec {tasks} tâches "
            "de projet et {recs} recommandations de ressources."
        ),
        "risk_summary_template": "{n} risques de couverture de tâches à examiner.",
    },
    "German": {
        "report_title": "Ressourcenzuordnungsbericht",
        "original_query": "Ursprüngliche Anfrage",
        "executive_summary": "Zusammenfassung",
        "how_to_read": "So lesen Sie diesen Bericht",
        "project_breakdown": "Projektübersicht",
        "task_recs": "Aufgabenbasierte Ressourcenempfehlungen",
        "risk_flags_title": "Risikohinweise",
        "task": "Aufgabe", "workstream": "Arbeitsstream",
        "skills_needed": "Benötigte Fähigkeiten", "strength": "Stärke",
        "effort": "Aufwand", "resources": "Ressourcen",
        "fit": "Eignung", "availability": "Verfügbarkeit", "rationale": "Begründung",
        "hrs_wk": "Std/Wo", "primary": "Primär", "support": "Unterstützung",
        "no_tasks": "Keine Projektaufgaben generiert.",
        "no_recs": "Keine Aufgabenempfehlungen generiert.",
        "needs_staffing": "Personaleinsatz überprüfen",
        "no_risks": "Keine Risikohinweise.",
        "risk_prefix": "Risiko",
        "no_problem": "Keine Problembeschreibung angegeben.",
        "strength_def": (
            "Stärke: aggregierte Teamkompetenz für eine Fähigkeit oder Aufgabe, "
            "basierend auf Berater-Kompetenzlevel-Gewichten und semantischer "
            "Ähnlichkeit mit den erforderlichen Fähigkeiten."
        ),
        "fit_def": (
            "Eignung: aufgabenbezogene Übereinstimmung zwischen Berater-Fähigkeiten, "
            "semantischer Ähnlichkeit und Verfügbarkeit."
        ),
        "summary_template": (
            "Aufgabenbasierter Personalplan mit {tasks} Projektaufgaben und "
            "{recs} Ressourcenempfehlungen erstellt."
        ),
        "risk_summary_template": "{n} Aufgaben-Abdeckungsrisiken müssen überprüft werden.",
    },
    "Spanish": {
        "report_title": "Informe de Correspondencia de Recursos",
        "original_query": "Consulta Original",
        "executive_summary": "Resumen Ejecutivo",
        "how_to_read": "Cómo Leer Este Informe",
        "project_breakdown": "Desglose del Proyecto",
        "task_recs": "Recomendaciones de Recursos por Tarea",
        "risk_flags_title": "Indicadores de Riesgo",
        "task": "Tarea", "workstream": "Flujo de Trabajo",
        "skills_needed": "Habilidades Requeridas", "strength": "Fortaleza",
        "effort": "Esfuerzo", "resources": "Recursos",
        "fit": "Ajuste", "availability": "Disponibilidad", "rationale": "Justificación",
        "hrs_wk": "hrs/sem", "primary": "Primario", "support": "Soporte",
        "no_tasks": "No se generaron tareas de proyecto.",
        "no_recs": "No se generaron recomendaciones de tareas.",
        "needs_staffing": "Requiere revisión de personal",
        "no_risks": "Sin indicadores de riesgo.",
        "risk_prefix": "Riesgo",
        "no_problem": "No se proporcionó descripción del problema.",
        "strength_def": (
            "Fortaleza: capacidad agregada del equipo para una habilidad o tarea, "
            "basada en pesos de nivel de habilidad de consultores y similitud "
            "semántica con las habilidades requeridas."
        ),
        "fit_def": (
            "Ajuste: correspondencia a nivel de tarea entre habilidades del consultor, "
            "similitud semántica y disponibilidad."
        ),
        "summary_template": (
            "Plan de dotación basado en tareas generado con {tasks} tareas de "
            "proyecto y {recs} recomendaciones de recursos."
        ),
        "risk_summary_template": "{n} riesgos de cobertura de tareas necesitan revisión.",
    },
    "Portuguese": {
        "report_title": "Relatório de Correspondência de Recursos",
        "original_query": "Consulta Original",
        "executive_summary": "Resumo Executivo",
        "how_to_read": "Como Ler Este Relatório",
        "project_breakdown": "Divisão do Projeto",
        "task_recs": "Recomendações de Recursos por Tarefa",
        "risk_flags_title": "Sinalizadores de Risco",
        "task": "Tarefa", "workstream": "Fluxo de Trabalho",
        "skills_needed": "Habilidades Necessárias", "strength": "Força",
        "effort": "Esforço", "resources": "Recursos",
        "fit": "Adequação", "availability": "Disponibilidade", "rationale": "Justificativa",
        "hrs_wk": "hrs/sem", "primary": "Primário", "support": "Suporte",
        "no_tasks": "Nenhuma tarefa de projeto gerada.",
        "no_recs": "Nenhuma recomendação de tarefa gerada.",
        "needs_staffing": "Necessita revisão de equipe",
        "no_risks": "Nenhum sinalizador de risco.",
        "risk_prefix": "Risco",
        "no_problem": "Nenhuma descrição de problema fornecida.",
        "strength_def": (
            "Força: capacidade agregada da equipe para uma habilidade ou tarefa, "
            "baseada em pesos de nível de habilidade de consultores e similaridade "
            "semântica com as habilidades necessárias."
        ),
        "fit_def": (
            "Adequação: correspondência no nível de tarefa entre habilidades do "
            "consultor, similaridade semântica e disponibilidade."
        ),
        "summary_template": (
            "Plano de alocação baseado em tarefas gerado com {tasks} tarefas de "
            "projeto e {recs} recomendações de recursos."
        ),
        "risk_summary_template": "{n} riscos de cobertura de tarefas precisam de revisão.",
    },
    "Arabic": {
        "report_title": "تقرير مطابقة الموارد",
        "original_query": "الاستعلام الأصلي",
        "executive_summary": "الملخص التنفيذي",
        "how_to_read": "كيفية قراءة هذا التقرير",
        "project_breakdown": "تفصيل المشروع",
        "task_recs": "توصيات موارد المهام",
        "risk_flags_title": "علامات المخاطر",
        "task": "مهمة", "workstream": "تدفق العمل",
        "skills_needed": "المهارات المطلوبة", "strength": "القوة",
        "effort": "الجهد", "resources": "الموارد",
        "fit": "الملاءمة", "availability": "التوفر", "rationale": "المبرر",
        "hrs_wk": "ساعة/أسبوع", "primary": "أساسي", "support": "داعم",
        "no_tasks": "لم يتم إنشاء مهام مشروع.",
        "no_recs": "لم يتم إنشاء توصيات مهام.",
        "needs_staffing": "يحتاج مراجعة التوظيف",
        "no_risks": "لا توجد علامات مخاطر.",
        "risk_prefix": "خطر",
        "no_problem": "لم يتم توفير وصف للمشكلة.",
        "strength_def": (
            "القوة: القدرة الإجمالية للفريق على مهارة أو مهمة، بناءً على أوزان "
            "مستوى مهارة الاستشاريين والتشابه الدلالي مع المهارات المطلوبة."
        ),
        "fit_def": (
            "الملاءمة: مطابقة على مستوى المهمة بين مهارات الاستشاري والتشابه "
            "الدلالي والتوفر."
        ),
        "summary_template": (
            "تم إنشاء خطة توظيف قائمة على المهام مع {tasks} مهمة مشروع "
            "و{recs} توصية موارد على مستوى المهمة."
        ),
        "risk_summary_template": "{n} مخاطر تغطية مهام تحتاج إلى مراجعة.",
    },
    "Japanese": {
        "report_title": "リソースマッチングレポート",
        "original_query": "元の質問",
        "executive_summary": "エグゼクティブサマリー",
        "how_to_read": "レポートの読み方",
        "project_breakdown": "プロジェクト概要",
        "task_recs": "タスクリソース推薦",
        "risk_flags_title": "リスクフラグ",
        "task": "タスク", "workstream": "ワークストリーム",
        "skills_needed": "必要なスキル", "strength": "強度",
        "effort": "工数", "resources": "リソース",
        "fit": "適合度", "availability": "稼働時間", "rationale": "理由",
        "hrs_wk": "時間/週", "primary": "主担当", "support": "サポート",
        "no_tasks": "プロジェクトタスクが生成されませんでした。",
        "no_recs": "タスク推薦が生成されませんでした。",
        "needs_staffing": "スタッフィングレビューが必要",
        "no_risks": "リスクフラグなし。",
        "risk_prefix": "リスク",
        "no_problem": "問題説明が提供されていません。",
        "strength_def": (
            "強度: コンサルタントのスキルレベルの重みと必要スキルとの意味的類似性に基づく、"
            "スキルまたはタスクに対するチームの総合能力。"
        ),
        "fit_def": (
            "適合度: コンサルタントのスキル、意味的類似性、および稼働時間に基づく"
            "タスクレベルのマッチング。"
        ),
        "summary_template": (
            "{tasks}プロジェクトタスクと{recs}タスクレベルリソース推薦を含む"
            "タスクベースのスタッフィングプランを生成しました。"
        ),
        "risk_summary_template": "{n}件のタスクカバレッジリスクの確認が必要です。",
    },
    "Chinese (Simplified)": {
        "report_title": "资源匹配报告",
        "original_query": "原始问题",
        "executive_summary": "执行摘要",
        "how_to_read": "如何阅读本报告",
        "project_breakdown": "项目分解",
        "task_recs": "任务资源建议",
        "risk_flags_title": "风险标志",
        "task": "任务", "workstream": "工作流",
        "skills_needed": "所需技能", "strength": "强度",
        "effort": "工作量", "resources": "资源",
        "fit": "匹配度", "availability": "可用时间", "rationale": "理由",
        "hrs_wk": "小时/周", "primary": "主要", "support": "支持",
        "no_tasks": "未生成项目任务。",
        "no_recs": "未生成任务建议。",
        "needs_staffing": "需要人员审查",
        "no_risks": "无风险标志。",
        "risk_prefix": "风险",
        "no_problem": "未提供问题描述。",
        "strength_def": (
            "强度：基于顾问技能水平权重和与任务所需技能的语义相似性，"
            "团队对某项技能或任务的综合能力。"
        ),
        "fit_def": (
            "匹配度：顾问技能、语义相似性和可用性之间的任务级别匹配。"
        ),
        "summary_template": (
            "已生成包含{tasks}个项目任务和{recs}个任务级别资源建议的"
            "基于任务的人员配置计划。"
        ),
        "risk_summary_template": "{n}个任务覆盖风险需要审查。",
    },
}


def _lb(language: str) -> dict:
    return _LABELS.get(language, _LABELS["English"])


def generate_report_bytes(data: dict, language: str = "English") -> bytes:
    """Generate a .docx report and return as bytes."""
    lb = _lb(language)
    doc = Document()
    _configure_document(doc)
    _add_title(doc, data.get("problem", ""), lb)
    _add_summary(doc, data, lb)
    _add_definitions(doc, lb)
    _add_project_breakdown(doc, data.get("solution", []),
                           data.get("task_recommendations", []), lb)
    _add_task_recommendations(doc, data.get("task_recommendations", []), lb)
    _add_risk_flags(doc, data.get("risk_flags", []), lb)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(31, 41, 55)


def _add_title(doc: Document, problem: str, lb: dict) -> None:
    title = doc.add_heading(lb["report_title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(lb["original_query"], level=1)
    doc.add_paragraph(problem or lb["no_problem"])


def _add_summary(doc: Document, data: dict, lb: dict) -> None:
    solution = data.get("solution", [])
    tasks = data.get("task_recommendations", [])
    risk_flags = data.get("risk_flags", [])
    doc.add_heading(lb["executive_summary"], level=1)
    assigned_count = sum(len(t.get("resources", [])) for t in tasks)
    doc.add_paragraph(
        lb["summary_template"].format(tasks=len(solution), recs=assigned_count)
    )
    if risk_flags:
        doc.add_paragraph(lb["risk_summary_template"].format(n=len(risk_flags)))


def _add_definitions(doc: Document, lb: dict) -> None:
    doc.add_heading(lb["how_to_read"], level=1)
    doc.add_paragraph(lb["strength_def"])
    doc.add_paragraph(lb["fit_def"])


def _add_project_breakdown(
    doc: Document, solution: list[dict], tasks: list[dict], lb: dict,
) -> None:
    doc.add_heading(lb["project_breakdown"], level=1)
    if not solution:
        doc.add_paragraph(lb["no_tasks"])
        return
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = [lb["task"], lb["workstream"], lb["skills_needed"],
                lb["strength"], lb["effort"]]
    _write_table_row(table.rows[0].cells, headers, bold=True)
    task_by_id = {t.get("task_id"): t for t in tasks}
    for step in solution:
        task = task_by_id.get(step.get("step"), {})
        required = task.get("required_skills", [])
        _write_table_row(table.add_row().cells, [
            str(step.get("step", "?")),
            step.get("action", "N/A"),
            ", ".join(required) or step.get("technology", "N/A"),
            str(task.get("strength_score", step.get("skill_strength_score", "N/A"))),
            step.get("effort", "TBD"),
        ])


def _add_task_recommendations(
    doc: Document, tasks: list[dict], lb: dict,
) -> None:
    doc.add_heading(lb["task_recs"], level=1)
    if not tasks:
        doc.add_paragraph(lb["no_recs"])
        return
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = [lb["task"], lb["skills_needed"], lb["resources"],
                lb["fit"], lb["availability"], lb["rationale"]]
    _write_table_row(table.rows[0].cells, headers, bold=True)
    rec_map = {"Primary": lb["primary"], "Support": lb["support"]}
    for task in tasks:
        resources = task.get("resources", [])
        resource_names = [
            f"{rec_map.get(r.get('recommendation', ''), r.get('recommendation', ''))}: "
            f"{r.get('full_name', 'N/A')}"
            for r in resources
        ]
        fits = [
            f"{r.get('full_name', 'N/A')}: {r.get('fit_score', 'N/A')}"
            for r in resources
        ]
        avail = [
            f"{r.get('full_name', 'N/A')}: {r.get('availability', 'N/A')} {lb['hrs_wk']}"
            for r in resources
        ]
        _write_table_row(table.add_row().cells, [
            task.get("task", "N/A"),
            ", ".join(task.get("required_skills", [])) or "N/A",
            "\n".join(resource_names) or lb["needs_staffing"],
            "\n".join(fits) or "N/A",
            "\n".join(avail) or "N/A",
            task.get("rationale", ""),
        ])


def _add_risk_flags(doc: Document, risk_flags: list[dict], lb: dict) -> None:
    doc.add_heading(lb["risk_flags_title"], level=1)
    if not risk_flags:
        doc.add_paragraph(lb["no_risks"])
        return
    for flag in risk_flags:
        doc.add_paragraph(
            f"{lb['risk_prefix']}: {flag.get('reason', 'Unknown risk')}",
            style="List Bullet",
        )


def _write_table_row(cells, values: list[str], bold: bool = False) -> None:
    for cell, value in zip(cells, values):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(str(value))
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(31, 41, 55)
